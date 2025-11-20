# main.py (updated)
import os
import io
import json
import hashlib
import re
import sqlite3
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime, timedelta

from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Request, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from jose import jwt, JWTError

import fitz  # PyMuPDF
import docx
from docx import Document as DocxDocument
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from base64 import b64encode, b64decode
from dotenv import load_dotenv

# AI models and helpers (unchanged)
from ai_models.indian_legal_bert import IndianLegalBERT
from ai_models.risk_engine import AdvancedRiskEngine
from utils.file_processor import FileProcessor

# LLM prediction helpers
from ai_models.llm_predictions import (
    llm_predict_clause_type,
    llm_generate_summary,
    llm_generate_improved_clause,
)

load_dotenv()

# -----------------------
# Config
# -----------------------
SECRET_KEY = os.getenv("SECRET_KEY", "dev-secret-change-me")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "60"))
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploads")
AES_KEY_B64 = os.getenv("AES_KEY_B64")
SINGLE_USER_EMAIL = os.getenv("SINGLE_USER_EMAIL", "admin@example.com")
SINGLE_USER_PASSWORD = os.getenv("SINGLE_USER_PASSWORD", "changeme")
SQLITE_DB = os.getenv("SQLITE_DB", "./data/app.db")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(os.path.dirname(SQLITE_DB) or ".", exist_ok=True)

print("DEBUG LLM KEY:", os.getenv("LLM_API_KEY"))
print("DEBUG LLM_API_KEY present:", bool(os.getenv("LLM_API_KEY")))

# -----------------------
# AES helpers (unchanged)
# -----------------------
def _ensure_aes_key() -> bytes:
    if AES_KEY_B64:
        try:
            key = b64decode(AES_KEY_B64)
            if len(key) in (16, 24, 32):
                return key
            raise ValueError("AES_KEY_B64 decoded length invalid")
        except Exception as e:
            raise RuntimeError(f"Invalid AES_KEY_B64: {e}")
    # fallback to SHA256 of SECRET_KEY (32 bytes)
    return hashlib.sha256(SECRET_KEY.encode()).digest()[:32]


AES_KEY = _ensure_aes_key()

def _pad(b: bytes) -> bytes:
    pad_len = AES.block_size - (len(b) % AES.block_size)
    return b + bytes([pad_len]) * pad_len

def _unpad(b: bytes) -> bytes:
    return b[:-b[-1]]

def encrypt_bytes(data: bytes) -> str:
    iv = get_random_bytes(16)
    cipher = AES.new(AES_KEY, AES.MODE_CBC, iv)
    ct = cipher.encrypt(_pad(data))
    return b64encode(iv + ct).decode("utf-8")

def decrypt_bytes(b64str: str) -> bytes:
    raw = b64decode(b64str)
    iv, ct = raw[:16], raw[16:]
    cipher = AES.new(AES_KEY, AES.MODE_CBC, iv)
    return _unpad(cipher.decrypt(ct))

# -----------------------
# DB helpers (SQLite)
# -----------------------
def get_db_conn():
    conn = sqlite3.connect(SQLITE_DB, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_conn()
    cur = conn.cursor()
    # users table: id, email (unique), password_hash, salt, created_at
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        salt TEXT NOT NULL,
        created_at TEXT NOT NULL
    );
    """)
    # conversations: id, user_email, stored_filename, title, created_at, analysis_json (TEXT)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS conversations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_email TEXT NOT NULL,
        stored_filename TEXT,
        title TEXT,
        created_at TEXT NOT NULL,
        analysis_json TEXT
    );
    """)
    conn.commit()
    conn.close()

def create_user_in_db(email: str, password: str):
    salt = hashlib.sha256(os.urandom(16)).hexdigest()
    pwd_hash = hashlib.sha256((salt + password).encode()).hexdigest()
    conn = get_db_conn()
    cur = conn.cursor()
    try:
        cur.execute("INSERT INTO users (email, password_hash, salt, created_at) VALUES (?, ?, ?, ?)",
                    (email.lower(), pwd_hash, salt, datetime.utcnow().isoformat()))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        raise ValueError("User already exists")
    conn.close()

def verify_user_password(email: str, password: str) -> bool:
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT password_hash, salt FROM users WHERE email = ?", (email.lower(),))
    row = cur.fetchone()
    conn.close()
    if not row:
        return False
    expected = row["password_hash"]
    salt = row["salt"]
    return hashlib.sha256((salt + password).encode()).hexdigest() == expected

def user_exists(email: str) -> bool:
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM users WHERE email = ?", (email.lower(),))
    exists = cur.fetchone() is not None
    conn.close()
    return exists

def create_conversation(user_email: str, stored_filename: Optional[str], title: Optional[str]) -> int:
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("INSERT INTO conversations (user_email, stored_filename, title, created_at) VALUES (?, ?, ?, ?)",
                (user_email.lower(), stored_filename, title or stored_filename or "", datetime.utcnow().isoformat()))
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id

def update_conversation_analysis_by_stored(stored_filename: str, user_email: str, analysis: dict):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("UPDATE conversations SET analysis_json = ? WHERE stored_filename = ? AND user_email = ?",
                (json.dumps(analysis), stored_filename, user_email.lower()))
    conn.commit()
    conn.close()

def get_user_conversations(user_email: str) -> List[Dict[str, Any]]:
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT id, stored_filename, title, created_at, analysis_json FROM conversations WHERE user_email = ? ORDER BY created_at DESC",
                (user_email.lower(),))
    rows = cur.fetchall()
    conn.close()
    out = []
    for r in rows:
        out.append({
            "id": r["id"],
            "stored_filename": r["stored_filename"],
            "title": r["title"],
            "created_at": r["created_at"],
            "analysis": json.loads(r["analysis_json"]) if r["analysis_json"] else None
        })
    return out

def get_conversation(user_email: str, conv_id: int) -> Optional[Dict[str, Any]]:
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT id, stored_filename, title, created_at, analysis_json FROM conversations WHERE user_email = ? AND id = ?",
                (user_email.lower(), conv_id))
    r = cur.fetchone()
    conn.close()
    if not r:
        return None
    return {
        "id": r["id"],
        "stored_filename": r["stored_filename"],
        "title": r["title"],
        "created_at": r["created_at"],
        "analysis": json.loads(r["analysis_json"]) if r["analysis_json"] else None
    }

# Initialize DB
init_db()

# If no users & SINGLE_USER env is provided, create the legacy single user in DB for compatibility
try:
    if not user_exists(SINGLE_USER_EMAIL):
        create_user_in_db(SINGLE_USER_EMAIL, SINGLE_USER_PASSWORD)
        print("[INIT] Created single user from env for compatibility:", SINGLE_USER_EMAIL)
except Exception as e:
    print("[INIT WARN] Could not create single user:", e)

# -----------------------
# Auth helpers
# -----------------------
class Token(BaseModel):
    access_token: str
    token_type: str

class LoginRequest(BaseModel):
    email: str
    password: str

def create_access_token(data: dict) -> str:
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode = {**data, "exp": expire}
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_user(request: Request):
    auth = request.headers.get("Authorization", "")
    if not auth or not auth.startswith("Bearer "):
        raise HTTPException(401, "Missing Authorization header")
    token = auth.split(" ", 1)[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        sub = payload.get("sub")
        if not sub:
            raise HTTPException(401, "Invalid token (no sub)")
        # ensure user exists
        if not user_exists(sub):
            raise HTTPException(401, "Invalid user")
        return {"email": sub}
    except JWTError:
        raise HTTPException(401, "Invalid token")

# -----------------------
# Models init (unchanged)
# -----------------------
try:
    legal_bert = IndianLegalBERT()
    risk_engine = AdvancedRiskEngine(db_session=None)
    print("[DEBUG] AI models initialized.")
except Exception as e:
    print("[WARN] Failed to initialize AI models at startup:", str(e))
    class _BrokenModel:
        def __getattr__(self, name):
            def _boom(*a, **k):
                raise RuntimeError(f"Model not available: attempted to call {name}. Original error: {e}")
            return _boom
    legal_bert = _BrokenModel()
    risk_engine = _BrokenModel()

# -----------------------
# FastAPI app
# -----------------------
app = FastAPI(title="Legal AI Document Analyzer")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------
# Text extraction helpers (unchanged)
# -----------------------
def extract_text_from_pdf_bytes(b: bytes) -> str:
    try:
        text = []
        with fitz.open(stream=b, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text("text"))
        return "\n".join(text)
    except Exception:
        return b.decode("utf-8", errors="replace") if isinstance(b, (bytes, bytearray)) else str(b)

def extract_text_from_docx_bytes(b: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(b))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return b.decode("utf-8", errors="replace") if isinstance(b, (bytes, bytearray)) else str(b)

# -----------------------
# Document analysis helpers (unchanged)
# -----------------------
def generate_improved_clause(clause_text: str, clause_type: str):
    t = clause_text.strip()
    safe = t

    safe = re.sub(r'unlimited liability', "liability capped to a reasonable amount", safe, flags=re.I)
    safe = re.sub(r'50% per month', "2% per month (reasonable statutory limit)", safe, flags=re.I)

    if safe != t:
        return safe.rstrip(".") + "."

    # BERT summarizer rewrite
    try:
        out = legal_bert.summarizer(
            f"Rewrite safely: {clause_text}",
            max_length=250,
            truncation=True
        )
        if isinstance(out, list) and out[0].get("generated_text"):
            return out[0]["generated_text"].strip()
    except:
        pass

    return (
        "Revised safer clause: Obligations must comply with applicable Indian laws. "
        "Liability is limited to direct damages only, reasonably capped, and timelines must "
        "be clearly defined (e.g., 30 days)."
    )

def build_corrected_document(clauses: List[str], analysis_results: List[Dict[str, Any]], threshold: float = 0.20) -> str:
    improved_map = {}
    for r in analysis_results:
        score = float(r.get("risk_score", 0.0))
        if score >= threshold and r.get("improved_clause"):
            improved_map[r["clause_text"].strip()] = r["improved_clause"]

    corrected_parts = []
    for original in clauses:
        stripped = original.strip()
        if stripped in improved_map:
            corrected_parts.append(improved_map[stripped])
        else:
            corrected_parts.append(original)
    return "\n\n".join(corrected_parts)

def build_corrected_docx_bytes(corrected_text: str) -> bytes:
    doc = DocxDocument()
    title = doc.add_heading("AI-Corrected Legal Document", level=1)
    title.alignment = 1

    doc.add_heading("Corrected Document", level=2)

    for para in corrected_text.split("\n"):
        doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
def clean_document_boilerplate(text: str) -> str:
    """
    Removes judicial headers, e-stamp metadata, footers, disclaimers,
    and extracts only the actual agreement body.
    """

    # Common noisy patterns in Indian legal PDFs
    boilerplate_patterns = [
        r"eSahayak\.io.*?\n",                  # e-Sahayak watermark
        r"Sample.*?\n",                        # sample page text
        r"Stamp Duty.*?\n",                    # stamp header
        r"Crove\.app.*?\n",                    # Crove watermark
        r"Government of.*?\n",                 # Gov headers
        r"Notary.*?\n",
        r"Digitally signed.*?\n",
        r"Page \d+ of \d+",                    # page numbers
        r"-----.*?-----",                      # separators
    ]

    # Remove each pattern
    cleaned = text
    for p in boilerplate_patterns:
        cleaned = re.sub(p, "", cleaned, flags=re.IGNORECASE | re.DOTALL)

    # Now detect start of a real agreement
    start_markers = [
        "THIS AGREEMENT",
        "THIS DEED",
        "RENT AGREEMENT",
        "MEMORANDUM OF UNDERSTANDING",
        "NOW THIS AGREEMENT",
        "NOW THIS DEED",
        "WITNESSETH",
        "TERMS AND CONDITIONS",
    ]

    # Find earliest hit
    start_index = None
    for marker in start_markers:
        loc = cleaned.upper().find(marker)
        if loc != -1:
            if start_index is None or loc < start_index:
                start_index = loc

    # If marker found → cut everything before it
    if start_index is not None:
        cleaned = cleaned[start_index:]

    return cleaned.strip()

async def _analyze_document_internal(payload: dict) -> dict:
    text = payload.get("text")
    stored_filename = payload.get("stored_filename")

    # decrypt if stored
    if not text and stored_filename:
        path = os.path.join(UPLOAD_DIR, stored_filename)
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail="stored file not found")
        enc = open(path, "r", encoding="utf-8").read()
        try:
            raw = decrypt_bytes(enc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to decrypt stored file: {e}")

        lower = path.lower()
        if lower.endswith(".docx.enc"):
            text = extract_text_from_docx_bytes(raw)
        else:
            text = extract_text_from_pdf_bytes(raw) or extract_text_from_docx_bytes(raw) or (raw.decode("utf-8", errors="replace") if isinstance(raw, (bytes, bytearray)) else str(raw))

    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="text or stored_filename required")

    # clauses segmentation etc. (exactly same as before)
    try:
        clauses = FileProcessor.preprocess_text(text)
    except Exception:
        clauses = [c.strip() for c in re.split(r'\n{2,}', text) if c.strip()]

    if not clauses:
        clauses = [text]

    results: List[Dict[str, Any]] = []
    total_risk = 0.0
    MAX_CLAUSES = 50

    for clause in clauses[:MAX_CLAUSES]:
        if len(clause.strip()) < 20:
            continue
        clause_type = "general"
        confidence = 0.0
        try:
            ct = llm_predict_clause_type(clause)
            if ct and "clause_type" in ct:
                clause_type = ct["clause_type"]
                confidence = float(ct.get("confidence", 0.85))
            else:
                clause_type, confidence = legal_bert.predict_clause_type(clause)
        except Exception:
            try:
                clause_type, confidence = legal_bert.predict_clause_type(clause)
            except Exception:
                clause_type, confidence = "general", 0.0

        try:
            risk_data = risk_engine.analyze_risk_with_statutes(clause, clause_type)
        except Exception:
            risk_data = {
                "risk_level": "low",
                "risk_score": 0.0,
                "violations": [],
                "compliance_issues": [],
                "legal_references": [],
                "pattern_violations": []
            }
        risk_score = float(risk_data.get("risk_score", 0.0))

        try:
            legal_refs = risk_engine.get_legal_references(clause, clause_type)
        except Exception:
            legal_refs = risk_data.get("legal_references", [])

        try:
            recs = risk_engine.generate_dynamic_recommendations(clause, risk_data, legal_refs)
        except Exception:
            recs = []

        try:
            summary = llm_generate_summary(clause)
        except Exception as e:
            print("[LLM SUMMARY] Error:", e)
            summary = None

        if not summary:
            try:
                summary = legal_bert.generate_dynamic_summary(clause)
            except Exception:
                summary = (clause[:300] + "...") if len(clause) > 300 else clause

        improved_clause = None
        is_risky = risk_score >= 0.20

        if is_risky:
            try:
                improved_clause = llm_generate_improved_clause(clause)
            except Exception as e:
                print("[LLM IMPROVE] Error:", e)
                improved_clause = None

            if not improved_clause:
                try:
                    improved_clause = generate_improved_clause(clause, clause_type)
                except Exception as e:
                    print("[BERT IMPROVE] Error:", e)
                    improved_clause = None

        if is_risky and not improved_clause:
            improved_clause = (
                "Revised safer clause: All obligations must comply with Indian contract "
                "law. Liability is limited to direct, reasonable damages, and all timelines "
                "must be clearly defined and fair."
            )

        results.append({
            "clause_text": clause,
            "clause_type": clause_type,
            "confidence": float(confidence),
            "risk_level": risk_data.get("risk_level", "low"),
            "risk_score": risk_score,
            "summary": summary,
            "violations": risk_data.get("violations", []),
            "compliance_issues": risk_data.get("compliance_issues", []),
            "recommendations": recs,
            "legal_references": legal_refs,
            "improved_clause": improved_clause
        })

        total_risk += risk_score

    # Aggregation (same as previous code)...
    scores = [r.get("risk_score", 0.0) for r in results] or [0.0]
    max_score = max(scores)
    avg_all = sum(scores) / len(scores) if scores else 0.0
    TOP_K = 5
    top_k_scores = sorted(scores, reverse=True)[:TOP_K]
    avg_top = sum(top_k_scores) / len(top_k_scores) if top_k_scores else 0.0

    overall_risk_score = (0.50 * max_score) + (0.35 * avg_top) + (0.15 * avg_all)
    overall_risk_score = max(0.0, min(1.0, overall_risk_score))

    if overall_risk_score >= 0.70:
        overall_risk_level = "High"
    elif overall_risk_score >= 0.20:
        overall_risk_level = "Medium"
    else:
        overall_risk_level = "Low"

    # risky filtering
    RISK_THRESHOLD = 0.2
    top_sorted = sorted(results, key=lambda x: x["risk_score"], reverse=True)
    risky_clauses = [r for r in top_sorted if r.get("risk_score", 0.0) >= RISK_THRESHOLD]

    if not risky_clauses:
        detailed_summary = (
            "✅ No significant risks detected.\n"
            "All clauses in this document have risk scores below 20%.\n"
            "The document appears legally safe and compliant.\n"
        )
    else:
        detailed_summary = ""

    recs_acc: List[str] = []
    law_set = set()
    for r in risky_clauses:
        for rec in r.get("recommendations", []):
            if isinstance(rec, str):
                if rec.strip().startswith("📘") or rec.strip().lower().startswith("reference"):
                    continue
                if rec.strip() and rec.strip() not in recs_acc:
                    recs_acc.append(rec.strip())
        for ref in r.get("legal_references", []):
            if isinstance(ref, dict):
                statute = ref.get("statute", "").strip()
                section = ref.get("section", "").strip()
                if statute:
                    if section and section not in statute:
                        law_set.add(f"{statute} – {section}")
                    else:
                        law_set.add(statute)
            elif isinstance(ref, str):
                law_set.add(ref.strip())

    if not recs_acc and not risky_clauses:
        recs_acc = [
            "✅ No significant risks detected.",
            "The document appears legally compliant and safe."
        ]

    final_recommendations = []
    for r in recs_acc:
        if r not in final_recommendations:
            final_recommendations.append(r)
    final_recommendations = final_recommendations[:20]

    relevant_laws = sorted(list(law_set))

    risky_summary = ""
    if risky_clauses:
        risky_summary += "⚠️ Top Risky Clauses\n(Only clauses with risk ≥ 20% are shown)\n\n"
        for idx, r in enumerate(risky_clauses[:10], start=1):
            risk_pct = round(r.get("risk_score", 0.0) * 100)
            risky_summary += f"{idx}. {r.get('clause_type','Unknown').upper()} — {r.get('risk_level','Unknown').upper()} ({risk_pct}%)\n{r.get('summary')}\n\n"

    risky_clauses_out = []
    for r in risky_clauses[:10]:
        risky_clauses_out.append({
            "clause_type": r.get("clause_type"),
            "risk_level": r.get("risk_level"),
            "risk_score": float(r.get("risk_score", 0.0)),
            "summary": r.get("summary"),
            "clause_text": r.get("clause_text"),
            "improved_clause": r.get("improved_clause"),
            "recommendations": r.get("recommendations", [])[:6],
            "legal_references": r.get("legal_references", [])[:6]
        })

    detected_risks = [r.get("clause_type", "unknown") for r in results]

    try:
        document_summary = llm_generate_summary(text) or legal_bert.generate_dynamic_summary(text, max_length=150)
    except Exception:
        document_summary = " ".join([r["summary"] for r in results[:5]])

    aggregation_debug = {
        "max_clause_score": round(max_score * 100),
        "avg_top_k_score": round(avg_top * 100),
        "avg_all_score": round(avg_all * 100),
        "combined_overall_score": round(overall_risk_score * 100),
    }

    corrected_document = build_corrected_document(clauses, results, threshold=0.20)

    return {
        "summary": document_summary,
        "overall_risk_score": round(overall_risk_score, 4),
        "overall_risk_level": overall_risk_level,
        "detailed_summary": detailed_summary,
        "risky_clauses": risky_clauses_out,
        "risky_summary": risky_summary,
        "detected_risks": detected_risks,
        "recommendations": final_recommendations,
        "relevant_laws": relevant_laws,
        "corrected_document": corrected_document,
        "count": len(results),
        "message": "⚠️ Risky clauses detected." if risky_clauses else "✅ No significant risks found.",
        "has_corrections": any(
            r.get("improved_clause") and r.get("improved_clause") != r.get("clause_text")
            for r in risky_clauses
        ),
        "original_clauses": clauses,
        "aggregation_debug": aggregation_debug
    }

# -----------------------
# Endpoints
# -----------------------
@app.get("/health")
def health():
    return {"status": "ok", "time": datetime.utcnow().isoformat()}

@app.post("/register")
def register(body: LoginRequest):
    email = body.email.lower()
    password = body.password
    if not email or not password:
        raise HTTPException(status_code=400, detail="Email and password required")
    if user_exists(email):
        raise HTTPException(status_code=400, detail="User already exists")
    try:
        create_user_in_db(email, password)
        return JSONResponse({"message": "User registered successfully"})
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/login", response_model=Token)
def login(body: LoginRequest):
    # prefer DB-backed users
    if verify_user_password(body.email, body.password):
        return {"access_token": create_access_token({"sub": body.email.lower()}), "token_type": "bearer"}
    # fallback to legacy single-user env if DB check failed
    if body.email == SINGLE_USER_EMAIL and body.password == SINGLE_USER_PASSWORD:
        # ensure user exists in db
        if not user_exists(body.email):
            create_user_in_db(body.email, body.password)
        return {"access_token": create_access_token({"sub": body.email}), "token_type": "bearer"}
    raise HTTPException(status_code=401, detail="Invalid credentials")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), user=Depends(get_current_user)):
    fname = file.filename
    content = await file.read()

    if not fname.lower().endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only .pdf and .docx supported")

    enc = encrypt_bytes(content)
    stamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    h = hashlib.sha256(content).hexdigest()[:12]
    out_name = f"{stamp}_{h}_{os.path.basename(fname)}.enc"
    out_path = os.path.join(UPLOAD_DIR, out_name)

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(enc)

    # create conversation row for this upload (title uses original filename)
    try:
        conv_id = create_conversation(user_email=user["email"], stored_filename=out_name, title=fname)
    except Exception as e:
        print("[WARN] Failed to create conversation row:", e)

    return {"stored_filename": out_name}

@app.post("/analyze")
async def analyze_document(payload: dict = Body(...), user=Depends(get_current_user)):
    try:
        result = await _analyze_document_internal(payload)
        # If stored_filename present, update conversation with analysis_json
        stored_filename = payload.get("stored_filename")
        if stored_filename:
            try:
                update_conversation_analysis_by_stored(stored_filename, user["email"], result)
            except Exception as e:
                print("[WARN] Failed to update conversation analysis:", e)
        return result
    except Exception as e:
        print("ERROR in /analyze:", e)
        raise HTTPException(status_code=500, detail=f"Internal error: {e}")

@app.post("/download-docx")
async def download_docx(payload: dict, user=Depends(get_current_user)):
    corrected_text = payload.get("corrected_text", "")
    filename = payload.get("filename", "Corrected_Document.docx")

    if not corrected_text:
        raise HTTPException(status_code=400, detail="Missing corrected_text")

    # create docx bytes and stream back
    bytes_data = build_corrected_docx_bytes(corrected_text)
    out_path = os.path.join(UPLOAD_DIR, filename)
    with open(out_path, "wb") as f:
        f.write(bytes_data)

    response = FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Expose-Headers"] = "Content-Disposition"
    return response

# New endpoints for conversations
@app.get("/conversations")
def list_conversations(user=Depends(get_current_user)):
    try:
        convs = get_user_conversations(user["email"])
        return {"conversations": convs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/conversation/{conv_id}")
def get_conversation_endpoint(conv_id: int, user=Depends(get_current_user)):
    conv = get_conversation(user["email"], conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conv

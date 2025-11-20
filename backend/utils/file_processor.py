import fitz  # PyMuPDF
import docx
import re
from typing import List, Dict
import logging

class FileProcessor:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)

    @staticmethod
    def get_pdf_page_count(file_path: str) -> int:
        """Get page count of PDF file"""
        doc = fitz.open(file_path)
        page_count = doc.page_count
        doc.close()
        return page_count

    @staticmethod
    def get_docx_page_count(file_path: str) -> int:
        """Get approximate page count of DOCX file"""
        doc = docx.Document(file_path)
        # Rough estimation: 50 lines per page
        total_lines = sum(1 for paragraph in doc.paragraphs if paragraph.text.strip())
        return max(1, total_lines // 50)

    @staticmethod
    def preprocess_text(text: str) -> List[str]:
        """Preprocess and segment legal text into clauses with improved segmentation"""
        # Remove excessive whitespace and normalize
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Improved segmentation using multiple indicators
        # Split by common legal clause indicators with context preservation
        clauses = []
        
        # First, try to split by section/clause/article indicators
        section_splits = re.split(r'(?=Section\s+\d+|Clause\s+\d+|Article\s+\d+|^\s*[A-Z][A-Z\s]*\s*:)', text)
        
        for section in section_splits:
            if len(section.strip()) < 50:  # Skip very short sections
                continue
                
            # Further split long sections by sentence boundaries if they're too long
            sentences = re.split(r'[.!?]+\s+', section)
            current_clause = ""
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                    
                # If adding this sentence would make the clause too long, start a new one
                if len(current_clause + sentence) > 1000 and current_clause:
                    clauses.append(current_clause.strip())
                    current_clause = sentence
                else:
                    current_clause += " " + sentence
            
            if current_clause.strip():
                clauses.append(current_clause.strip())
        
        # Filter out clauses that are too short
        clauses = [clause for clause in clauses if len(clause.strip()) > 50]
        
        return clauses

    @staticmethod
    def extract_risk_keywords(clause: str) -> list:
        """Extract potentially risky keywords from clause"""
        risk_keywords = [
            # High-risk terms
            'penalty', 'fine', 'forfeit', 'terminate', 'default', 'liability',
            'unlimited', 'irrevocable', 'exclusive', 'non-compete', 'bond',
            'indemnify', 'waive', 'forfeit', 'forfeiture', 'automatic',
            'compulsory', 'mandatory', 'irreversible', 'irrevocable',
            'notwithstanding', 'despite', 'however', 'nevertheless',
            
            # Ambiguous terms
            'reasonable', 'appropriate', 'satisfactory', 'adequate', 'proper',
            'necessary', 'sufficient', 'timely', 'prompt', 'due',
            'at discretion', 'as required', 'if necessary', 'when possible',
            
            # Negative terms
            'not', 'never', 'no', 'except', 'unless', 'without',
            'prohibited', 'restricted', 'limited', 'subject to',
            'notwithstanding', 'despite', 'however', 'except when',
            
            # Conditional terms
            'if', 'provided', 'assuming', 'contingent', 'dependent',
            'conditional', 'subject to', 'pending', 'upon condition',
            
            # Legal jargon
            'hereby', 'whereas', 'therefore', 'notwithstanding',
            'pursuant to', 'in accordance with', 'as per'
        ]
        
        found_keywords = []
        clause_lower = clause.lower()
        
        for keyword in risk_keywords:
            if keyword in clause_lower:
                found_keywords.append(keyword)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_keywords = []
        for kw in found_keywords:
            if kw not in seen:
                seen.add(kw)
                unique_keywords.append(kw)
        
        return unique_keywords

    @staticmethod
    def analyze_document_structure(text: str) -> Dict:
        """Analyze document structure to provide better segmentation context"""
        analysis = {
            'total_length': len(text),
            'paragraph_count': len(text.split('\n\n')),
            'word_count': len(text.split()),
            'sentence_count': len(re.split(r'[.!?]+', text)),
            'clause_indicators': [],
            'section_indicators': [],
            'recommended_split_points': []
        }
        
        # Find clause indicators
        clause_pattern = r'(Section|Clause|Article|Part|Chapter)\s+\d+'
        clause_matches = re.finditer(clause_pattern, text, re.IGNORECASE)
        analysis['clause_indicators'] = [match.group() for match in clause_matches]
        
        # Find section headers
        section_pattern = r'^\s*[A-Z][A-Z\s]*\s*:|^.*\d+\.\s+.*$'
        section_matches = re.finditer(section_pattern, text, re.MULTILINE)
        analysis['section_indicators'] = [match.group() for match in section_matches]
        
        return analysis